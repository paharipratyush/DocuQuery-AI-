import streamlit as st
import faiss
import os
import validators
import traceback
import time
import threading
import logging
from typing import List, Union, Optional
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from streamlit.runtime.uploaded_file_manager import UploadedFile
from tqdm import tqdm
from google.api_core.exceptions import ResourceExhausted
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import re

# Set up logging
logging.basicConfig(
    filename="extraction_debug.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Initialize environment and configurations
try:
    from secret_api_keys import gemini_api_key
except ImportError:
    st.error("Error: secret_api_keys.py not found. Please create this file with your gemini_api_key.")
    st.stop()

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
CHUNK_SIZE = 500  # Original setting
CHUNK_OVERLAP = 100
URL_TIMEOUT = 30  # seconds
MAX_OUTPUT_TOKENS = 300  # Original setting
REQUESTS_PER_MINUTE = 15  # Free-tier limit for gemini-2.0-flash
MAX_CONTINUATIONS = 4
MAX_DIVE_DEEPER = 2

# Token Bucket for rate limiting
class TokenBucket:
    def __init__(self, tokens_per_minute=REQUESTS_PER_MINUTE):
        self.rate = tokens_per_minute / 60.0
        self.capacity = tokens_per_minute
        self.tokens = tokens_per_minute
        self.last_refill = time.time()
        self.lock = threading.Lock()

    def get_token(self):
        with self.lock:
            now = time.time()
            elapsed = now - self.last_refill
            new_tokens = elapsed * self.rate
            self.tokens = min(self.capacity, self.tokens + new_tokens)
            self.last_refill = now
            if self.tokens >= 1:
                self.tokens -= 1
                return True
            return False

# Global token bucket
token_bucket = TokenBucket()

class DocumentProcessor:
    """Handles document processing for various input types."""

    def __init__(self):
        self.embeddings = self._initialize_embeddings()

    @staticmethod
    @st.cache_resource
    def _initialize_embeddings():
        """Initialize and return the embedding model."""
        model_name = "sentence-transformers/all-mpnet-base-v2"
        return HuggingFaceEmbeddings(
            model_name=model_name,
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )

    def validate_url(self, url: str) -> bool:
        """Validate if the given URL is properly formatted."""
        return validators.url(url)

    def process_urls(self, urls: List[str]) -> List[str]:
        """Process URLs using WebBaseLoader for blogs and Selenium for academic faculty pages."""
        texts = []
        seen_urls = set()
        driver = None

        try:
            for url in tqdm(urls, desc="Processing URLs"):
                if url.strip() and url not in seen_urls and self.validate_url(url):
                    seen_urls.add(url)
                    text = ""
                    # Enhanced academic URL detection
                    is_academic = any(keyword in url.lower() for keyword in [
                        "faculty", "university", "department", "staff", "professor", "academic", "college", "institute"
                    ])

                    # Use WebBaseLoader for non-academic URLs (blogs)
                    if not is_academic:
                        try:
                            loader = WebBaseLoader(url)
                            loader.requests_kwargs = {
                                "timeout": URL_TIMEOUT,
                                "headers": {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"}
                            }
                            documents = loader.load()
                            extracted_texts = [doc.page_content for doc in documents if doc.page_content.strip()]
                            text = "\n".join(extracted_texts)
                            logging.info(f"Extracted with WebBaseLoader from {url}: {text[:200]}...")
                        except Exception as e:
                            logging.warning(f"WebBaseLoader failed for {url}: {str(e)}")
                            st.warning(f"Error processing URL {url}: {str(e)}")
                            text = ""

                    # Use Selenium for academic URLs or if WebBaseLoader fails
                    if is_academic or not text:
                        if not driver:
                            chrome_options = Options()
                            chrome_options.add_argument("--headless")
                            chrome_options.add_argument("--disable-gpu")
                            chrome_options.add_argument("--no-sandbox")
                            chrome_options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36")
                            driver = webdriver.Chrome(
                                service=Service(ChromeDriverManager().install()),
                                options=chrome_options
                            )
                        try:
                            driver.get(url)
                            time.sleep(10)  # Increased wait for dynamic content
                            # Enhanced scrolling
                            for _ in range(3):  # Multiple scrolls
                                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
                                time.sleep(3)
                            soup = BeautifulSoup(driver.page_source, "html.parser")
                            text = soup.get_text(separator="\n", strip=True)

                            if "cloudflare" in text.lower() or "please verify you are not a robot" in text.lower():
                                logging.warning(f"Cloudflare detected for {url}")
                                st.warning(f"Cloudflare security check detected for {url}. Try pasting the content using 'Text' input.")
                                continue

                            # Enhanced faculty extraction
                            if is_academic:
                                faculty_data = []
                                # Broader name pattern
                                name_pattern = re.compile(r"^(?:Dr\.|Prof\.|Mr\.|Ms\.|Mrs\.|[A-Z][a-z]+)?\s*[A-Z][a-z]+(?:\s[A-Z][a-z]+)*(?:,\s*(PhD|Ph.D.|M.Tech|MSc|MA|B.Tech))?$")
                                for element in soup.find_all(["div", "li", "tr", "p", "span", "h1", "h2", "h3", "h4"]):
                                    element_text = element.get_text(strip=True)
                                    if name_pattern.match(element_text) and len(element_text.split()) >= 2:
                                        # Flexible designation search
                                        designation = None
                                        for tag in ["p", "span", "div", "td", "li"]:
                                            next_elements = element.find_all_next(tag, limit=3)
                                            for next_elem in next_elements:
                                                next_text = next_elem.get_text(strip=True).lower()
                                                if any(keyword in next_text for keyword in [
                                                    "professor", "hod", "head", "faculty", "lecturer", "instructor", "dean", "chair"
                                                ]):
                                                    designation = next_elem.get_text(strip=True)
                                                    break
                                            if designation:
                                                break
                                        if not designation:
                                            parent = element.find_parent(["div", "li", "tr"])
                                            if parent:
                                                next_sibling = parent.find_next_sibling(["div", "li", "tr"])
                                                if next_sibling:
                                                    designation = next_sibling.get_text(strip=True)
                                        faculty_data.append(f"{element_text}: {designation or 'Faculty Member'}")
                                # Validate faculty list
                                if faculty_data:
                                    faculty_text = "\nFaculty List:\n" + "\n".join(sorted(set(faculty_data)))  # Remove duplicates
                                    text += faculty_text
                                    logging.info(f"Faculty List from {url}:\n{faculty_text}")
                                else:
                                    logging.warning(f"No faculty data extracted from {url}")
                                    st.warning(f"No faculty data extracted from {url}. The page may have an unusual structure.")

                        except Exception as e:
                            logging.error(f"Error processing URL {url} with Selenium: {str(e)}")
                            st.warning(f"Error processing URL {url}: {str(e)}")
                            continue

                    if text.strip():
                        texts.append(text)
                    else:
                        logging.warning(f"No content extracted from {url}")
                        st.warning(f"No content extracted from {url}.")

                elif url in seen_urls:
                    st.warning(f"Duplicate URL skipped: {url}")
                elif url.strip() and not self.validate_url(url):
                    st.warning(f"Invalid URL skipped: {url}")

        finally:
            if driver:
                driver.quit()

        if not texts:
            st.error("No content could be extracted from the provided URLs. Check extraction_debug.log for details.")
        return texts

    def process_pdf(self, file: Union[BytesIO, UploadedFile]) -> str:
        """Process PDF file and extract text."""
        if isinstance(file, UploadedFile) and file.size > MAX_FILE_SIZE:
            raise ValueError(f"File size exceeds limit of {MAX_FILE_SIZE/1024/1024}MB")

        try:
            if isinstance(file, UploadedFile):
                pdf_reader = PdfReader(BytesIO(file.read()))
            else:
                pdf_reader = PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            raise ValueError(f"Error processing PDF: {str(e)}")

    def process_docx(self, file: Union[BytesIO, UploadedFile]) -> str:
        """Process DOCX file and extract text."""
        try:
            if isinstance(file, UploadedFile):
                doc = Document(BytesIO(file.read()))
            else:
                doc = Document(file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            raise ValueError(f"Error processing DOCX: {str(e)}")

    def process_txt(self, file: Union[BytesIO, UploadedFile]) -> str:
        """Process TXT file and extract text."""
        try:
            if isinstance(file, UploadedFile):
                return file.read().decode("utf-8")
            else:
                return file.read().decode("utf-8")
        except Exception as e:
            raise ValueError(f"Error processing TXT: {str(e)}")

    def create_vectorstore(self, texts: List[str], existing_vectorstore: Optional[FAISS] = None) -> FAISS:
        """Create or update a FAISS vectorstore from the given texts."""
        try:
            if not texts or all(not text.strip() for text in texts):
                raise ValueError("No valid text content provided to create vector store.")

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=CHUNK_SIZE,
                chunk_overlap=CHUNK_OVERLAP,
                length_function=len,
                separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
            )

            split_texts = text_splitter.split_text(" ".join(texts))

            if not split_texts:
                raise ValueError("No text chunks generated after splitting.")

            # Prioritize faculty list chunks
            prioritized_texts = []
            for text in split_texts:
                if "Faculty List:" in text:
                    prioritized_texts.insert(0, text)  # Place at start
                else:
                    prioritized_texts.append(text)

            if existing_vectorstore:
                existing_vectorstore.add_texts(prioritized_texts)
                return existing_vectorstore
            else:
                dimension = len(self.embeddings.embed_query("test"))
                index = faiss.IndexFlatIP(dimension)
                vector_store = FAISS(
                    embedding_function=self.embeddings.embed_query,
                    index=index,
                    docstore=InMemoryDocstore(),
                    index_to_docstore_id={}
                )
                vector_store.add_texts(prioritized_texts)
                return vector_store

        except Exception as e:
            raise ValueError(f"Error creating vector store: {str(e)}")

def reset_session_state():
    """Reset all session state variables."""
    for key in list(st.session_state.keys()):
        del st.session_state[key]

def answer_question(vectorstore: FAISS, query: str, continuation: bool = False, continuation_count: int = 0, dive_deeper: bool = False) -> dict:
    """Generate answer for the given query using the vectorstore."""
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash",
            google_api_key=gemini_api_key,
            temperature=0.3,
            max_output_tokens=MAX_OUTPUT_TOKENS,
        )

        # Enhanced prompt for faculty and blog queries
        prompt_template = """You are an expert software engineer specializing in IoT and backend development, with the ability to extract specific details from academic and institutional contexts or summarize blog content. Answer the question concisely using the provided context. Structure the response based on the question type:

- If the question asks to "summarize" or "summary":
  - Provide a 2-3 sentence summary in simple, non-technical language first.
  - Follow with 2-3 bullet points of key details with moderate technical depth.
- For technical questions (e.g., about code, implementation):
  - Provide 2-3 bullet points of key concepts with clear explanations.
  - Include one short code snippet (if relevant).
  - Add a brief implementation note with practical insights.
- For specific fact questions (e.g., "Who is X?", "What is the name of Y?"):
  - Extract and provide the exact detail (e.g., name, role, number) from the context, prioritizing the "Faculty List" section if available.
  - If the detail is not found, state: "The specific information was not found in the context. Try rephrasing or checking the source."
  - Keep the answer concise, focusing on the requested fact.
- For list questions (e.g., "List all X", "What are all Y?"):
  - Provide a complete list of all relevant items (e.g., names, roles) from the "Faculty List" section in the context.
  - If the list is long, include as many items as possible within the response limit and end with "Continued in next part" to indicate more items.
  - If no "Faculty List" is found, state: "No faculty list was found in the context. Try rephrasing or checking the source."
- If "dive deeper" is requested:
  - Focus on advanced, intricate details (e.g., edge cases, optimizations, alternative approaches).
  - Provide 2-3 bullet points with in-depth technical insights.
  - Include a code snippet or example if applicable.
- Limit the answer to essential information relevant to the question and context.
- Initial responses should be moderately detailed; continuations should add deeper insights or continue lists.
- If the answer is complete, end with "Answer complete."
- If no more details are available, end with "No more details available."
- If the answer exceeds the response limit, end with "Continued in next part" and stop.
- For continuation requests (max {max_continuations} parts), resume from the last point with additional details or list items, avoiding repetition. For dive-deeper requests, focus on advanced insights.

Context: {context}

Question: {question}

Answer: """
        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"],
            template_format="f-string",
            partial_variables={"max_continuations": MAX_CONTINUATIONS}
        )

        qa = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type="stuff",
            retriever=vectorstore.as_retriever(
                search_kwargs={"k": 7}
            ),
            chain_type_kwargs={"prompt": prompt}
        )

        if dive_deeper and "last_query" in st.session_state:
            query = f"Provide advanced, intricate details about: {st.session_state['last_query']}"
        elif continuation and "last_query" in st.session_state and "last_answer" in st.session_state:
            query = f"Continue the answer to: {st.session_state['last_query']} from: {st.session_state['last_answer'][-100:]} with additional details or list items."

        # Rate limiting with token bucket
        max_attempts = 5
        for attempt in range(max_attempts):
            if token_bucket.get_token():
                try:
                    return qa({"query": query})
                except ResourceExhausted as e:
                    if "429" in str(e):
                        delay = 2 ** attempt * 2
                        if attempt < max_attempts - 1:
                            st.warning(f"Rate limit hit. Retrying in {delay} seconds...")
                            logging.info(f"429 error on attempt {attempt + 1}, waiting {delay}s")
                            time.sleep(delay)
                        else:
                            raise ValueError(f"Error generating answer: {str(e)}")
            else:
                delay = 60 / REQUESTS_PER_MINUTE
                st.warning(f"Rate limit approaching. Please wait {int(delay)} seconds before retrying.")
                logging.info(f"No tokens available, waiting {delay}s")
                time.sleep(delay)
        raise ValueError("Failed to acquire token after maximum attempts.")

    except Exception as e:
        raise ValueError(f"Error generating answer: {str(e)}")

def main():
    st.set_page_config(page_title="DocuQuery AI", layout="wide")
    st.title("📚 DocuQuery AI")

    # Initialize processor
    doc_processor = DocumentProcessor()

    # Sidebar
    with st.sidebar:
        st.header("Settings")
        if st.button("Reset Session"):
            reset_session_state()
            st.success("Session reset successfully!")
        debug_mode = st.checkbox("Enable Debug Mode (Show Extracted Content)", help="Displays raw extracted text to verify content or troubleshoot issues.")

    # Main input selection
    input_type = st.selectbox("Select Input Type", ["Link", "PDF", "Text", "DOCX", "TXT"])

    try:
        input_data = None
        if input_type == "Link":
            number_input = st.number_input(
                "Number of URLs",
                min_value=1,
                max_value=20,
                value=1,
                step=1
            )
            input_data = []
            for i in range(int(number_input)):
                url = st.text_input(f"URL {i+1}", key=f"url_{i}")
                if url:
                    input_data.append(url)

        elif input_type == "Text":
            input_data = st.text_area("Enter Text", height=200)
            if input_data and len(input_data.strip()) < 10:
                st.error("Text input is too short. Please provide at least 10 characters.")
                input_data = None

        elif input_type in ["PDF", "DOCX", "TXT"]:
            file_type = {"PDF": ["pdf"], "DOCX": ["docx", "doc"], "TXT": ["txt"]}[input_type]
            uploaded_file = st.file_uploader(f"Upload {input_type} file", type=file_type)
            if uploaded_file:
                if uploaded_file.size > MAX_FILE_SIZE:
                    st.error(f"File size exceeds maximum limit of {MAX_FILE_SIZE/1024/1024}MB")
                else:
                    input_data = uploaded_file

        # Process button
        if st.button("Proceed", disabled=not input_data):
            try:
                with st.spinner("Processing input..."):
                    if input_type == "Link":
                        texts = doc_processor.process_urls(input_data)
                    elif input_type == "PDF":
                        texts = [doc_processor.process_pdf(input_data)]
                    elif input_type == "DOCX":
                        texts = [doc_processor.process_docx(input_data)]
                    elif input_type == "TXT":
                        texts = [doc_processor.process_txt(input_data)]
                    else:
                        texts = [input_data]

                    if not texts or all(not text.strip() for text in texts):
                        st.error("No valid content extracted from the input. Check extraction_debug.log for details.")
                    else:
                        vectorstore = doc_processor.create_vectorstore(
                            texts,
                            existing_vectorstore=st.session_state.get("vectorstore")
                        )
                        st.session_state["vectorstore"] = vectorstore
                        st.session_state["extracted_texts"] = texts
                        st.success("Processing complete! You can now ask questions.")
                        if debug_mode:
                            st.subheader("DEBUG: Extracted Content")
                            for i, text in enumerate(texts):
                                st.write(f"**Source {i+1}:**")
                                st.text_area(f"Content {i+1}", text, height=200, key=f"debug_text_{i}")
                                if len(text) < 200:
                                    st.warning(f"Extracted content for Source {i+1} is short. Check if the page was fully loaded or use 'Text' input.")
                                if "Faculty List:" in text and len([line for line in text.split("\n") if ":" in line and "Faculty List" not in line]) < 2:
                                    st.warning(f"Faculty list for Source {i+1} is incomplete. The page may have an unusual structure.")
                                if st.button(f"Copy Content {i+1}", key=f"copy_button_{i}"):
                                    st.write("Content copied to clipboard! Select and copy manually.")

            except Exception as e:
                st.error(f"Error during processing: {str(e)}")
                st.error(traceback.format_exc())

        # Question answering section
        if "vectorstore" in st.session_state:
            st.header("Ask Questions")
            query = st.text_input("Enter your question:", key="query_input")

            if st.button("Get Answer", disabled=not query):
                try:
                    with st.spinner("Generating answer..."):
                        st.session_state["continuation_count"] = 0
                        st.session_state["dive_deeper_count"] = 0
                        answer = answer_question(st.session_state["vectorstore"], query)
                        st.session_state["last_query"] = query
                        st.session_state["last_answer"] = answer["result"]
                        st.write("Answer:", answer["result"])
                        is_truncated = (
                            len(answer["result"]) >= MAX_OUTPUT_TOKENS - 20 and
                            "Answer complete" not in answer["result"] and
                            "No more details available" not in answer["result"]
                        )
                        is_list_query = "list all" in query.lower() or "what are all" in query.lower()
                        st.session_state["is_truncated"] = is_truncated or (is_list_query and "Continued in next part" in answer["result"])
                        logging.info(f"Truncation detected: {is_truncated}, List query: {is_list_query}, Answer length: {len(answer['result'])}")
                        if st.session_state["is_truncated"]:
                            st.info("More details or list items available! Click 'Continue Answer' to read more.")
                        elif "Answer complete" in answer["result"]:
                            st.success("All essential information provided! Use 'Dive Deeper' for advanced insights.")
                        elif "No more details available" in answer["result"]:
                            st.success("No further details available for this question.")

                except Exception as e:
                    st.error(f"Error generating answer: {str(e)}")

            # Continue Answer button
            if st.session_state.get("is_truncated", False) and st.session_state.get("continuation_count", 0) < MAX_CONTINUATIONS:
                if st.button("Continue Answer"):
                    try:
                        with st.spinner("Continuing answer..."):
                            st.session_state["continuation_count"] = st.session_state.get("continuation_count", 0) + 1
                            answer = answer_question(
                                st.session_state["vectorstore"],
                                query,
                                continuation=True,
                                continuation_count=st.session_state["continuation_count"]
                            )
                            st.session_state["last_answer"] += "\n\n" + answer["result"]
                            st.write("Continued Answer:", answer["result"])
                            is_truncated = (
                                len(answer["result"]) >= MAX_OUTPUT_TOKENS - 20 and
                                "Answer complete" not in answer["result"] and
                                "No more details available" not in answer["result"]
                            )
                            is_list_query = "list all" in query.lower() or "what are all" in query.lower()
                            st.session_state["is_truncated"] = is_truncated or (is_list_query and "Continued in next part" in answer["result"])
                            logging.info(f"Continuation truncation detected: {is_truncated}, List query: {is_list_query}, Answer length: {len(answer['result'])}, Continuation count: {st.session_state['continuation_count']}")
                            if st.session_state["is_truncated"] and st.session_state["continuation_count"] < MAX_CONTINUATIONS:
                                st.info("More details or list items available! Click 'Continue Answer' to read more.")
                            elif st.session_state["continuation_count"] >= MAX_CONTINUATIONS:
                                st.warning("Maximum continuation limit reached. Use 'Dive Deeper' for advanced insights or refine your question.")
                                st.session_state["is_truncated"] = False
                            elif "Answer complete" in answer["result"]:
                                st.success("All essential information provided! Use 'Dive Deeper' for advanced insights.")
                                st.session_state["is_truncated"] = False
                            elif "No more details available" in answer["result"]:
                                st.success("No further details available for this question.")
                                st.session_state["is_truncated"] = False
                    except Exception as e:
                        st.error(f"Error continuing answer: {str(e)}")

            # Dive Deeper button
            if "last_answer" in st.session_state and st.session_state.get("dive_deeper_count", 0) < MAX_DIVE_DEEPER:
                if st.button("Dive Deeper"):
                    try:
                        with st.spinner("Fetching advanced details..."):
                            st.session_state["dive_deeper_count"] = st.session_state.get("dive_deeper_count", 0) + 1
                            answer = answer_question(
                                st.session_state["vectorstore"],
                                query,
                                dive_deeper=True,
                                continuation_count=st.session_state["continuation_count"]
                            )
                            st.session_state["last_answer"] += "\n\n" + answer["result"]
                            st.write("Advanced Details:", answer["result"])
                            is_truncated = (
                                len(answer["result"]) >= MAX_OUTPUT_TOKENS - 20 and
                                "Answer complete" not in answer["result"] and
                                "No more details available" not in answer["result"]
                            )
                            st.session_state["is_truncated"] = is_truncated
                            logging.info(f"Dive deeper detected: {is_truncated}, Answer length: {len(answer['result'])}, Dive deeper count: {st.session_state['dive_deeper_count']}")
                            if is_truncated and st.session_state["dive_deeper_count"] < MAX_DIVE_DEEPER:
                                st.info("More advanced details available! Click 'Dive Deeper' again.")
                            elif st.session_state["dive_deeper_count"] >= MAX_DIVE_DEEPER:
                                st.warning("Maximum dive-deeper limit reached. Refine your question for more insights.")
                            elif "Answer complete" in answer["result"]:
                                st.success("All advanced details provided!")
                            elif "No more details available" in answer["result"]:
                                st.success("No further advanced details available.")
                    except Exception as e:
                        st.error(f"Error fetching advanced details: {str(e)}")

    except Exception as e:
        st.error(f"An unexpected error occurred: {str(e)}")
        st.error(traceback.format_exc())

if __name__ == "__main__":
    main()
